import pandas as pd
import numpy as np
import os
from num2words import num2words
from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx2pdf import convert

def create_offer_first(calc_name):
    #set doc name
    docx_name_list = calc_name.split(sep = ".")
    docx_name_list = docx_name_list[:-1]
    docx_name_str = ""
    for elem in docx_name_list:
        docx_name_str = docx_name_str + elem

    docx_name = docx_name_str + '_first' + ".docx"

    #reading file
    excel = pd.read_excel(calc_name, "Калькуляция")
    print(excel)

    #get
    object = excel.iloc[search_obj(excel)][0]
    sum = round(excel.iloc[search_sum(excel)][4], 2)
    sum_rub = str(sum).split(sep = ".")[0]
    #split rub
    sum_rub_format = '{0:,}'.format(int(sum_rub)).replace(',', ' ')
    sum_cent = str(sum).split(sep = ".")[1]
    sum_context = sum_rub_format + ", " + sum_cent
    # sum to text
    sum_text = num2words(sum_rub, lang="ru") + " руб. " + num2words(sum_cent, lang = "ru") + " коп."
    sum_text_arr = sum_text.split(" ")
    if sum_text_arr[-2] == "один":
        sum_text_arr[-2] = "одна"
    if sum_text_arr[-2] == "два":
        sum_text_arr[-2] = "две"
    sum_text = ""
    for elem in sum_text_arr:
        sum_text = sum_text + elem + " "
    sum_text = sum_text[:-1:]
    print(object)
    print(sum)
    print(sum_text)

    # вставить название объекта и стоимость
    doc = DocxTemplate("Шаблон ТКП.docx")
    context = { 'object' : object, 'sum' : sum_context, 'sum_text': sum_text}
    doc.render(context)

    # сохранение ТКП 
    doc.save("ТКП1.docx")

    # вставить таблицу спецификации
    # создание документа
    doc = Document("ТКП1.docx")
    # данные таблицы без названий колонок
    #reading file
    try:
        excel = pd.read_excel(calc_name, "Материалы").fillna(" ")
    except:
        excel = pd.read_excel(calc_name, "Осн.оборудование").fillna(" ")
    print(excel)

    #analyze and format
    start = search_first(excel)
    end = search_last(excel)
    data = excel.iloc[start:end, 0:4]
    data.columns = ["num", "art", "name", "count"]
    data.set_index(np.arange(0, len(data), 1), inplace=True)
    
    print(data)

    #collection
    items = tuple(data.itertuples(index = False, name = None))
    print(items)
    # добавляем таблицу материалов с одной строкой 
    # для заполнения названий колонок
    table = doc.add_table(1, len(items[0]))
    col = table.columns[0] 

    # стиль таблицы
    table.style = 'Table Grid'
    table.autofit = False

    # Получаем строку с колонками из добавленной таблицы 
    head_cells = table.rows[0].cells
    # добавляем названия колонок
    for i, item in enumerate(['№ п/п', 'Артикул', 'Наименование', 'Кол-во']):
        p = head_cells[i].paragraphs[0]
        # название колонки
        p.add_run(item).bold = True
        # выравниваем посередине
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        head_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # добавляем данные к существующей таблице
    for row in items:
        # добавляем строку с ячейками к объекту таблицы
        cells = table.add_row().cells
        for i, item in enumerate(row):
            # вставляем данные в ячейки
            cells[i].text = str(item)
            # изменим шрифт
            cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # ширина столбцов
    width = [0.5, 2, 3, 1]          
    for i, column in enumerate(table.columns):
        for cell in table.columns[i].cells:
            cell.width = Inches(width[i])

    # удалить лишний файл
    os.remove("ТКП1.docx")
    # сохранить файл предложения
    doc.save(docx_name)
    convert(docx_name, docx_name.split(sep = ".")[0] + ".pdf")

def create_offer_second(calc_name):
    #set doc name
    docx_name_list = calc_name.split(sep = ".")
    docx_name_list = docx_name_list[:-1]
    docx_name_str = ""
    for elem in docx_name_list:
        docx_name_str = docx_name_str + elem

    docx_name = docx_name_str + '_second' + ".docx"

    #reading file
    excel = pd.read_excel(calc_name, "Калькуляция")
    print(excel)

    #get
    object = excel.iloc[search_obj(excel)][0]
    sum = round(excel.iloc[search_sum(excel)][4], 2)
    sum_rub = str(sum).split(sep = ".")[0]
    #split rub
    sum_rub_format = '{0:,}'.format(int(sum_rub)).replace(',', ' ')
    sum_cent = str(sum).split(sep = ".")[1]
    sum_context = sum_rub_format + ", " + sum_cent
    # sum to text
    sum_text = num2words(sum_rub, lang="ru") + " руб. " + num2words(sum_cent, lang = "ru") + " коп."
    sum_text_arr = sum_text.split(" ")
    if sum_text_arr[-2] == "один":
        sum_text_arr[-2] = "одна"
    if sum_text_arr[-2] == "два":
        sum_text_arr[-2] = "две"
    sum_text = ""
    for elem in sum_text_arr:
        sum_text = sum_text + elem + " "
    sum_text = sum_text[:-1:]
    print(object)
    print(sum)
    print(sum_text)

    # вставить название объекта и стоимость
    doc = DocxTemplate("Шаблон ТКП.docx")
    context = { 'object' : object, 'sum' : sum_context, 'sum_text': sum_text}
    doc.render(context)

    # сохранение ТКП 
    doc.save("ТКП1.docx")

    # вставить таблицу спецификации
    # создание документа
    doc = Document("ТКП1.docx")
    # данные таблицы без названий колонок
    #reading file
    try:
        excel = pd.read_excel(calc_name, "Материалы").fillna(" ")
    except:
        excel = pd.read_excel(calc_name, "Осн.оборудование").fillna(" ")
    print(excel)

    #analyze and format
    start = search_first(excel)
    end = search_last(excel)
    data = excel.iloc[start:end, 0:6]
    data.columns = ["num", "art", "name", "count", "price", "sum"]
    data.set_index(np.arange(0, len(data), 1), inplace=True)
    data["price"] = data["price"].astype(float)
    data["sum"] = data["sum"].astype(float)
    data["price"] = data["price"].map('{:.2f}'.format)
    data["sum"] = data["sum"].map('{:.2f}'.format)
    data["price"] = data["price"].astype("string")
    data["sum"] = data["sum"].astype("string")
    data["price"] = data["price"].str.replace('.',',')
    data["sum"] = data["sum"].str.replace('.',',')
    
    print(data)

    #collection
    items = tuple(data.itertuples(index = False, name = None))
    print(items)
    # добавляем таблицу материалов с одной строкой 
    # для заполнения названий колонок
    table = doc.add_table(1, len(items[0]))
    col = table.columns[0] 

    # стиль таблицы
    table.style = 'Table Grid'
    table.autofit = False

    # Получаем строку с колонками из добавленной таблицы 
    head_cells = table.rows[0].cells
    # добавляем названия колонок
    for i, item in enumerate(['№ п/п', 'Артикул', 'Наименование', 'Кол-во', 'Цена руб., без НДС', 'Стоимость руб., без НДС']):
        p = head_cells[i].paragraphs[0]
        # название колонки
        p.add_run(item).bold = True
        # выравниваем посередине
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        head_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # добавляем данные к существующей таблице
    for row in items:
        # добавляем строку с ячейками к объекту таблицы
        cells = table.add_row().cells
        for i, item in enumerate(row):
            # вставляем данные в ячейки
            cells[i].text = str(item)
            # изменим шрифт
            cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # ширина столбцов
    width = [0.5, 1, 2.5, 0.5, 1, 1]          
    for i, column in enumerate(table.columns):
        for cell in table.columns[i].cells:
            cell.width = Inches(width[i])

    # удалить лишний файл
    os.remove("ТКП1.docx")
    # сохранить файл предложения
    doc.save(docx_name)
    convert(docx_name, docx_name.split(sep = ".")[0] + ".pdf")

def create_offer_third(calc_name):
    #set doc name
    docx_name_list = calc_name.split(sep = ".")
    docx_name_list = docx_name_list[:-1]
    docx_name_str = ""
    for elem in docx_name_list:
        docx_name_str = docx_name_str + elem

    docx_name = docx_name_str + ".docx"

    #reading file
    excel = pd.read_excel(calc_name, "Калькуляция")
    print(excel)

    #get
    object = excel.iloc[search_obj(excel)][0]
    sum = round(excel.iloc[search_sum(excel)][4], 2)
    sum_rub = str(sum).split(sep = ".")[0]
    #split rub
    sum_rub_format = '{0:,}'.format(int(sum_rub)).replace(',', ' ')
    sum_cent = str(sum).split(sep = ".")[1]
    sum_context = sum_rub_format + ", " + sum_cent
    # sum to text
    sum_text = num2words(sum_rub, lang="ru") + " руб. " + num2words(sum_cent, lang = "ru") + " коп."
    sum_text_arr = sum_text.split(" ")
    if sum_text_arr[-2] == "один":
        sum_text_arr[-2] = "одна"
    if sum_text_arr[-2] == "два":
        sum_text_arr[-2] = "две"
    sum_text = ""
    for elem in sum_text_arr:
        sum_text = sum_text + elem + " "
    sum_text = sum_text[:-1:]
    print(object)
    print(sum)
    print(sum_text)

    # вставить название объекта и стоимость
    doc = DocxTemplate("Шаблон ТКП.docx")
    context = { 'object' : object, 'sum' : sum_context, 'sum_text': sum_text}
    doc.render(context)

    # сохранение ТКП 
    doc.save("ТКП1.docx")

    # вставить таблицу спецификации
    # создание документа
    doc = Document("ТКП1.docx")
    # данные таблицы без названий колонок
    #reading file
    try:
        excel = pd.read_excel(calc_name, "Материалы").fillna(" ")
    except:
        excel = pd.read_excel(calc_name, "Осн.оборудование").fillna(" ")
    print(excel)

    #analyze and format
    start = search_first(excel)
    end = search_last(excel)
    data = excel.iloc[start:end, 0:6]
    data.columns = ["num", "art", "name", "count", "price", "sum"]
    data.set_index(np.arange(0, len(data), 1), inplace=True)
    data["price"] = data["price"].astype(float)
    data["sum"] = data["sum"].astype(float)
    data["price"] = data["price"].map('{:.2f}'.format)
    data["sum"] = data["sum"].map('{:.2f}'.format)
    data["price"] = data["price"].astype("string")
    data["sum"] = data["sum"].astype("string")
    data["price"] = data["price"].str.replace('.',',')
    data["sum"] = data["sum"].str.replace('.',',')
    
    print(data)

    #collection
    items = tuple(data.itertuples(index = False, name = None))
    print(items)
    # добавляем таблицу материалов с одной строкой 
    # для заполнения названий колонок
    table = doc.add_table(1, len(items[0]))
    col = table.columns[0] 

    # стиль таблицы
    table.style = 'Table Grid'
    table.autofit = False

    # Получаем строку с колонками из добавленной таблицы 
    head_cells = table.rows[0].cells
    # добавляем названия колонок
    for i, item in enumerate(['№ п/п', 'Артикул', 'Наименование', 'Кол-во', 'Цена руб., без НДС', 'Стоимость руб., без НДС']):
        p = head_cells[i].paragraphs[0]
        # название колонки
        p.add_run(item).bold = True
        # выравниваем посередине
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        head_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # добавляем данные к существующей таблице
    for row in items:
        # добавляем строку с ячейками к объекту таблицы
        cells = table.add_row().cells
        for i, item in enumerate(row):
            # вставляем данные в ячейки
            cells[i].text = str(item)
            # изменим шрифт
            cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # ширина столбцов
    width = [0.5, 1, 2.5, 0.5, 1, 1]          
    for i, column in enumerate(table.columns):
        for cell in table.columns[i].cells:
            cell.width = Inches(width[i])

    #переход к новой таблице
    empty_paragraph = doc.add_paragraph(" ")
    empty_paragraph.style.font.size = Pt(1)

    #additional collection
    excel = pd.read_excel(calc_name, "Калькуляция")
    materials_sum, work_sum, fee_sum, total_sum = get_additional(excel)
    materials_sum = '{:.2f}'.format(materials_sum)
    work_sum = '{:.2f}'.format(work_sum)
    fee_sum = '{:.2f}'.format(fee_sum)
    total_sum = '{:.2f}'.format(total_sum)
    additional_items = [tuple(["Стоимость изготовления:", str(work_sum).replace(".", ",")]), tuple(["НДС 20%", str(fee_sum).replace(".", ",")]), tuple(["Стоимость изделия с НДС", str(total_sum).replace(".", ",")])]
    print(additional_items)
    # добавляем таблицу с одной строкой 
    # для заполнения названий колонок
    additional_table = doc.add_table(1, len(additional_items[0]))
    #col = additional_table.columns[0] 

    # стиль таблицы
    additional_table.style = 'Table Grid'
    additional_table.autofit = False

    head_cells = additional_table.rows[0].cells
    # добавляем названия колонок
    for i, item in enumerate(["Итого стоимость комплектующих:", str(materials_sum).replace(".", ",")]):
        p = head_cells[i].paragraphs[0]
        # название колонки
        p.add_run(item).bold = False
        head_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # добавляем данные к существующей таблице
    for row in additional_items:
        # добавляем строку с ячейками к объекту таблицы
        cells = additional_table.add_row().cells
        for i, item in enumerate(row):
            # вставляем данные в ячейки
            cells[i].text = str(item)
            # изменим шрифт
            cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # ширина столбцов
    additional_width = [5.5, 1]           
    for i, column in enumerate(additional_table.columns):
        for cell in additional_table.columns[i].cells:
            cell.width = Inches(additional_width[i])

    # удалить лишний файл
    os.remove("ТКП1.docx")
    # сохранить файл предложения
    doc.save(docx_name)
    # PDF
    convert(docx_name, docx_name.split(sep = ".")[0] + ".pdf")
    # подготовка excel
    # list of dataframes
    data.columns = ['№ п/п', 'Артикул', 'Наименование', 'Кол-во', 'Цена руб., без НДС', 'Стоимость руб., без НДС']
    data.set_index('№ п/п', inplace=True)
    additional_data = pd.DataFrame([["Итого стоимость комплектующих:", str(materials_sum).replace(".", ",")], ["Стоимость изготовления:", str(work_sum).replace(".", ",")], ["НДС 20%", str(fee_sum).replace(".", ",")], ["Стоимость изделия с НДС", str(total_sum).replace(".", ",")]], columns = ["Итого стоимость комплектующих:", str(materials_sum).replace(".", ",")])
    additional_data = additional_data[1:]
    additional_data.set_index(additional_data.iloc[:,0], inplace=True)
    dfs = [data, additional_data.iloc[:,1:]]

    # excel
    multiple_dfs(dfs, 'Материалы', "Эксперт " + docx_name_str + ".xlsx", 0)


    
# funtion
def multiple_dfs(df_list, sheets, file_name, spaces):
    writer = pd.ExcelWriter(file_name,engine='xlsxwriter')   
    row = 0
    for dataframe in df_list:
        dataframe.to_excel(writer,sheet_name=sheets,startrow=row , startcol=0)   
        row = row + len(dataframe.index) + spaces + 1
    writer.close()

    

# search object
def search_obj(excel_file):
    for i in range(len(excel_file)):
        if excel_file.iloc[i][0] == "Калькуляция на изготовление" or excel_file.iloc[i][0] == "Калькуляция на работы по изготовлению":
            return i + 1
        
# search sum
def search_sum(excel_file):
    for i in range(len(excel_file)):
        if excel_file.iloc[i][0] == "ВСЕГО:":
            return i
        
# search first pos
def search_first(excel_file):
    for i in range(len(excel_file)):
        if excel_file.iloc[i][0] == "№ п/п":
            return i+1
        
# search last pos
def search_last(excel_file):
    for i in range(len(excel_file)):
        if excel_file.iloc[i][0] == "ИТОГО:":
            return i

#search additional info
def get_additional(excel_file):
    materials_sum = 0
    work_sum = 0
    fee_sum = 0
    total_sum = 0

    for i in range(len(excel_file)):
        if excel_file.iloc[i][1] == "Основное оборудование" or excel_file.iloc[i][1] == "Основные материалы":
            materials_sum = round(float(excel_file.iloc[i][4]), 2)
            print(materials_sum)

    for i in range(len(excel_file)):
        if excel_file.iloc[i][0] == "ИТОГО:":
            work_sum = round(float(excel_file.iloc[i][4]) - materials_sum, 2)
            print(work_sum)

    for i in range(len(excel_file)):
        if excel_file.iloc[i][0] == "НДС - 20%":
            fee_sum = round(float(excel_file.iloc[i][4]), 2)
            print(fee_sum)
    
    for i in range(len(excel_file)):
        if excel_file.iloc[i][0] == "ВСЕГО:":
            total_sum = round(float(excel_file.iloc[i][4]), 2)
            print(total_sum)
        
    return materials_sum, work_sum, fee_sum, total_sum
    
    

# choose file
from tkinter.filedialog import askopenfilenames
initdir=os.getcwd()
calculations = []
calculations = askopenfilenames(initialdir=initdir, title="Choose calc files")
calculations_names = []
for calc in calculations:
    calc = calc.split(sep = "/")[-1]
    calculations_names.append(calc)
print(calculations_names)

# main cycle
for calc in calculations_names:
    create_offer_first(calc)
    create_offer_second(calc)
    create_offer_third(calc)