import pandas as pd
import numpy as np
import os
from num2words import num2words
from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

def create_offer(calc_name):
    #set doc name
    docx_name = calc_name.split(sep = ".")[0] + ".docx"

    #reading file
    excel = pd.read_excel(calc_name, "Калькуляция")
    print(excel)

    #get
    object = excel.iloc[search_obj(excel)][0]
    sum = round(excel.iloc[search_sum(excel)][4], 2)
    sum_rub = str(sum).split(sep = ".")[0]
    #split rub
    sum_rub_format = '{0:,}'.format(int(sum_rub)).replace(',', ' ')
    sum_cent = str(sum).split(sep = ".")[1]
    sum_context = sum_rub_format + ", " + sum_cent
    # sum to text
    sum_text = num2words(sum_rub, lang="ru") + " руб. " + num2words(sum_cent, lang = "ru") + " коп."
    sum_text_arr = sum_text.split(" ")
    if sum_text_arr[-2] == "один":
        sum_text_arr[-2] = "одна"
    if sum_text_arr[-2] == "два":
        sum_text_arr[-2] = "две"
    sum_text = ""
    for elem in sum_text_arr:
        sum_text = sum_text + elem + " "
    sum_text = sum_text[:-1:]
    print(object)
    print(sum)
    print(sum_text)

    # вставить название объекта и стоимость
    doc = DocxTemplate("Шаблон ТКП.docx")
    context = { 'object' : object, 'sum' : sum_context, 'sum_text': sum_text}
    doc.render(context)
    doc.save("ТКП1.docx")

    # вставить таблицу спецификации
    # создание пустого документа
    doc = Document("ТКП1.docx")
    # данные таблицы без названий колонок
    #reading file
    excel = pd.read_excel(calc_name, "Материалы").fillna(" ")
    print(excel)

    #analyze and format
    start = search_first(excel)
    end = search_last(excel)
    data = excel.iloc[start:end, 0:4]
    data.columns = ["num", "art", "name", "count"]
    data.set_index(np.arange(0, len(data), 1), inplace=True)

    print(data)

    #collection
    items = tuple(data.itertuples(index = False, name = None))
    print(items)
    # добавляем таблицу с одной строкой 
    # для заполнения названий колонок
    table = doc.add_table(1, len(items[0]))
    col = table.columns[0] 

    # стиль таблицы
    table.style = 'Table Grid'
    table.autofit = False

    # Получаем строку с колонками из добавленной таблицы 
    head_cells = table.rows[0].cells
    # добавляем названия колонок
    for i, item in enumerate(['№ п/п', 'Артикул', 'Наименование', 'Кол-во']):
        p = head_cells[i].paragraphs[0]
        # название колонки
        p.add_run(item).bold = True
        # выравниваем посередине
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # добавляем данные к существующей таблице
    for row in items:
        # добавляем строку с ячейками к объекту таблицы
        cells = table.add_row().cells
        for i, item in enumerate(row):
            # вставляем данные в ячейки
            cells[i].text = str(item)
            # если последняя ячейка
            if i == 2:
                # изменим шрифт
                cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'

    # ширина столбцов
    width = [0.5, 2, 3.5, 0.5]          
    for i, column in enumerate(table.columns):
        for cell in table.columns[i].cells:
            cell.width = Inches(width[i])

    # удалить лишний файл
    os.remove("ТКП1.docx")
    # сохранить файл предложения
    doc.save(docx_name)

# search object
def search_obj(excel_file):
    for i in range(len(excel_file)):
        if excel_file.iloc[i][0] == "Калькуляция на изготовление":
            return i + 1
        
# search sum
def search_sum(excel_file):
    for i in range(len(excel_file)):
        if excel_file.iloc[i][0] == "ВСЕГО:":
            return i
        
# search first pos
def search_first(excel_file):
    for i in range(len(excel_file)):
        if excel_file.iloc[i][0] == "№ п/п":
            return i+1
        
# search last pos
def search_last(excel_file):
    for i in range(len(excel_file)):
        if excel_file.iloc[i][0] == "ИТОГО:":
            return i

# choose file
from tkinter.filedialog import askopenfilenames
initdir=os.getcwd()
calculations = []
calculations = askopenfilenames(initialdir=initdir, title="Choose calc files")
calculations_names = []
for calc in calculations:
    calc = calc.split(sep = "/")[-1]
    calculations_names.append(calc)
print(calculations_names)

# main cycle
for calc in calculations_names:
    create_offer(calc)